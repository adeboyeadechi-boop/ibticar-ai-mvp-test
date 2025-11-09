import os
from docx import Document
import json

def extract_text_from_docx(file_path):
    """Extract all text from a DOCX file."""
    try:
        doc = Document(file_path)
        text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        return "\n".join(text)
    except Exception as e:
        return f"Error reading {file_path}: {str(e)}"

def scan_implementations_folder(base_path):
    """Scan all DOCX files in implementations folder."""
    results = {}

    for root, dirs, files in os.walk(base_path):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                file_path = os.path.join(root, file)
                relative_path = os.path.relpath(file_path, base_path)
                print(f"Processing: {relative_path}")

                content = extract_text_from_docx(file_path)
                results[relative_path] = content

    return results

if __name__ == "__main__":
    base_path = r"C:\Users\Adeboye Adechi\Documents\ADSERVIO\IBTICAR_AI\ibticar-ai-mvp\implementations"

    print("Starting extraction...")
    all_contents = scan_implementations_folder(base_path)

    # Save to JSON for analysis
    output_file = r"C:\Users\Adeboye Adechi\Documents\ADSERVIO\IBTICAR_AI\ibticar-ai-mvp\implementations_extracted.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_contents, f, ensure_ascii=False, indent=2)

    print(f"\nExtraction complete! Found {len(all_contents)} files.")
    print(f"Results saved to: {output_file}")

    # Print summary
    print("\nFiles processed:")
    for file_path in sorted(all_contents.keys()):
        content_length = len(all_contents[file_path])
        print(f"  - {file_path} ({content_length} chars)")
